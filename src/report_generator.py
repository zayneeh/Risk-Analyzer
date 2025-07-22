from docx import Document
from docx.shared import Inches, RGBColor
import matplotlib.pyplot as plt
import os

def generate_report(analyzed_data, output_path="outputs/rfe_risk_report.docx", extra_notes=None):
    doc = Document()
    doc.add_heading("EB-1A RFE Risk Report", level=0)
    doc.add_paragraph("This memo provides an automated risk review of the EB-1A petition based on USCIS adjudication standards.")

    # === Risk Timeline Chart ===
    doc.add_heading("📊 Risk Timeline", level=1)
    chart_path = "outputs/rfe_risk_timeline.png"
    create_risk_chart(analyzed_data, chart_path)
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))
        doc.add_paragraph("This chart visualizes the relative risk level of each section.")

    # === Extra Notes (Heuristic Results) ===
    if extra_notes:
        doc.add_heading("🔍 Heuristic Pattern Warnings", level=1)

        # Letter repetition warning
        sim_pairs = extra_notes.get("similar_letters", [])
        if sim_pairs:
            doc.add_paragraph("🟠 The following recommendation letters appear overly similar:")
            for a, b, score in sim_pairs:
                doc.add_paragraph(f"• {a} ↔ {b} (Similarity: {score})", style="List Bullet")

        # Field inconsistency warning
        conflicting = extra_notes.get("conflicting_fields", [])
        if conflicting:
            doc.add_paragraph("⚠️ Multiple fields of expertise were mentioned:")
            for field in conflicting:
                doc.add_paragraph(f"• {field}", style="List Bullet")

    # === Section-by-Section Review ===
    doc.add_heading("📋 Section-by-Section Findings", level=1)

    for item in analyzed_data:
        doc.add_heading(f"{item['criteria']} ({item['section']})", level=2)

        doc.add_paragraph("📄 Excerpt:\n" + item['excerpt'])

        doc.add_paragraph("🔎 Risk Analysis:")
        for bullet in item['llm_feedback'].split("\n"):
            if bullet.strip():
                doc.add_paragraph(bullet.strip("–• "), style='List Bullet')

        if item.get("buzzwords"):
            buzz = ", ".join(item["buzzwords"])
            doc.add_paragraph("⚠️ Buzzwords Flagged: " + buzz)

        if item.get("reviewer_voice"):
            para = doc.add_paragraph("🧑‍⚖️ USCIS Reviewer Note: ")
            run = para.add_run(item["reviewer_voice"])
            run.italic = True
            run.font.color.rgb = RGBColor(0, 102, 204)

    doc.save(output_path)


def create_risk_chart(analyzed_data, output_path="outputs/rfe_risk_timeline.png"):
    def infer_score(feedback):
        text = feedback.lower()
        if any(word in text for word in ["unsupported", "lacks", "generic", "vague"]):
            return "High"
        elif any(word in text for word in ["moderate", "could improve"]):
            return "Medium"
        return "Low"

    labels = [item["section"] for item in analyzed_data]
    risk_labels = [infer_score(item["llm_feedback"]) for item in analyzed_data]
    risk_map = {"Low": 1, "Medium": 2, "High": 3}
    scores = [risk_map[label] for label in risk_labels]

    colors = ["green" if s == 1 else "orange" if s == 2 else "red" for s in scores]

    plt.figure(figsize=(10, 4))
    plt.bar(labels, scores, color=colors)
    plt.xticks(rotation=30)
    plt.yticks([1, 2, 3], ["Low", "Medium", "High"])
    plt.title("📊 RFE Risk Timeline Across Petition Sections")
    plt.xlabel("Petition Section")
    plt.ylabel("Risk Level")
    plt.tight_layout()

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    plt.savefig(output_path)
    plt.close()

